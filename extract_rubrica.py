from docx import Document
from pathlib import Path
import sys


def main():
    p = Path(r"c:\vscode\cursos\6semestre\series de tiempo\dashboard_final\RÚBRICA DE EVALUACIÓN_dashboard.docx")
    try:
        doc = Document(p)
    except Exception as e:
        print("ERROR al abrir el documento:", e)
        sys.exit(1)

    print("=== PARÁRAFOS ===")
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            print(t)

    print("\n=== TABLAS ===")
    for i, table in enumerate(doc.tables, start=1):
        print(f"-- Tabla {i} --")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            print(" | ".join(cells))


if __name__ == "__main__":
    main()